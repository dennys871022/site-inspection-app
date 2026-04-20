import streamlit as st
try:
    import pkg_resources
except ImportError:
    import setuptools

from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docxcompose.composer import Composer
from PIL import Image
import io
import datetime
from datetime import timedelta, timezone
import os
import zipfile
import pandas as pd
import smtplib
import re
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication

# ==========================================
# 0. 雲端資料庫設定
# ==========================================
GOOGLE_SHEETS_CSV_URL = "https://docs.google.com/spreadsheets/d/1ubR0wOJkOhA4IYyQ_Qq-LUldKwkEj084N45Ym04sKU8/export?format=csv"

# ==========================================
# 1. 核心功能函式庫
# ==========================================

def get_taiwan_date():
    utc_now = datetime.datetime.now(timezone.utc)
    return (utc_now + timedelta(hours=8)).date()

def get_paragraph_style(paragraph):
    style = {}
    if paragraph.runs:
        run = paragraph.runs[0]
        style['font_name'] = run.font.name
        style['font_size'] = run.font.size
        style['bold'] = run.bold
        style['italic'] = run.italic
        style['underline'] = run.underline
        style['color'] = run.font.color.rgb
        try:
            rPr = run._element.rPr
            if rPr is not None and rPr.rFonts is not None:
                style['eastAsia'] = rPr.rFonts.get(qn('w:eastAsia'))
        except: pass
    return style

def apply_style_to_run(run, style):
    if not style: return
    if style.get('font_name'): run.font.name = style.get('font_name')
    if style.get('font_size'): run.font.size = style['font_size']
    if style.get('bold') is not None: run.bold = style['bold']
    if style.get('italic') is not None: run.italic = style['italic']
    if style.get('underline') is not None: run.underline = style['underline']
    if style.get('color'): run.font.color.rgb = style['color']
    if style.get('eastAsia'):
        run._element.rPr.rFonts.set(qn('w:eastAsia'), style['eastAsia'])
    elif style.get('font_name') == 'Times New Roman':
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

def compress_image(image_file, max_width=800):
    img = Image.open(image_file)
    if img.mode == 'RGBA': img = img.convert('RGB')
    try:
        from PIL import ImageOps
        img = ImageOps.exif_transpose(img)
    except: pass
    ratio = max_width / float(img.size[0])
    if ratio < 1:
        h_size = int((float(img.size[1]) * float(ratio)))
        img = img.resize((max_width, h_size), Image.Resampling.LANCZOS)
    img_byte_arr = io.BytesIO()
    img.save(img_byte_arr, format='JPEG', quality=75)
    img_byte_arr.seek(0)
    return img_byte_arr

def replace_text_content(doc, replacements):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_paragraph_pure(paragraph, replacements)
    for paragraph in doc.paragraphs:
        replace_paragraph_pure(paragraph, replacements)

def replace_paragraph_pure(paragraph, replacements):
    if not paragraph.text: return
    original_text = paragraph.text
    needs_replace = False
    for key in replacements:
        if key in original_text:
            needs_replace = True
            break
    if needs_replace:
        saved_style = get_paragraph_style(paragraph)
        new_text = original_text
        for key, value in replacements.items():
            val_str = str(value) if value is not None else ""
            new_text = new_text.replace(key, val_str)
        paragraph.clear()
        new_run = paragraph.add_run(new_text)
        apply_style_to_run(new_run, saved_style)

def replace_placeholder_with_image(doc, placeholder, image_stream):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        align = paragraph.alignment
                        paragraph.clear()
                        paragraph.alignment = align
                        run = paragraph.add_run()
                        if image_stream:
                            run.add_picture(image_stream, width=Cm(8.0))
                        return

def remove_element(element):
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)

def truncate_doc_after_page_break(doc):
    body = doc.element.body
    break_index = -1
    for i, element in enumerate(body):
        if element.tag.endswith('p'):
            if 'w:br' in element.xml and 'type="page"' in element.xml:
                break_index = i
                break
    if break_index != -1:
        for i in range(len(body) - 1, break_index - 1, -1):
            if body[i].tag.endswith('sectPr'):
                continue
            remove_element(body[i])

def generate_single_page(template_bytes, context, photo_batch, start_no):
    doc = Document(io.BytesIO(template_bytes))
    text_replacements = {f"{{{k}}}": v for k, v in context.items()}
    replace_text_content(doc, text_replacements)
    
    for i in range(1, 9):
        img_key = f"{{img_{i}}}"
        info_key = f"{{info_{i}}}"
        idx = i - 1
        if idx < len(photo_batch):
            data = photo_batch[idx]
            replace_placeholder_with_image(doc, img_key, compress_image(data['file']))
            
            spacer = "\u3000" * 4 
            
            info_text = f"照片編號：{data['no']:02d}{spacer}日期：{data['date_str']}\n"
            info_text += f"說明：{data['desc']}\n"
            
            if data.get('design') and data['design'].strip():
                info_text += f"設計：{data['design']}\n"
                
            info_text += f"實測：{data['result']}"
            
            replace_text_content(doc, {info_key: info_text})
        else:
            pass 

    if len(photo_batch) <= 4:
        truncate_doc_after_page_break(doc)
    
    final_clean = {}
    for i in range(1, 9):
        final_clean[f"{{img_{i}}}"] = ""
        final_clean[f"{{info_{i}}}"] = ""
    replace_text_content(doc, final_clean)

    return doc

def generate_names(selected_type, base_date):
    clean_type = selected_type.split(' (EA')[0].split(' (EB')[0]
    suffix = "自主檢查"
    if "施工" in clean_type or "混凝土" in clean_type:
        suffix = "施工自主檢查"
        clean_type = clean_type.replace("-施工", "")
    elif "材料" in clean_type:
        suffix = "材料進場自主檢查"
        clean_type = clean_type.replace("-材料", "")
    elif "有價廢料" in clean_type:
        suffix = "有價廢料清運自主檢查"
        clean_type = clean_type.replace("-有價廢料", "")
    
    match = re.search(r'(\(.*\))', clean_type)
    extra_info = ""
    if match:
        extra_info = match.group(1) 
        clean_type = clean_type.replace(extra_info, "").strip() 
        
    full_item_name = f"{clean_type}{suffix}{extra_info}"
    
    roc_year = base_date.year - 1911
    roc_date_str = f"{roc_year}{base_date.month:02d}{base_date.day:02d}"
    file_name = f"{roc_date_str}{full_item_name}"
    return full_item_name, file_name

def generate_clean_filename_base(selected_type, base_date):
    _, file_name = generate_names(selected_type, base_date)
    return file_name

def send_email_via_secrets(doc_bytes, filename, receiver_email, receiver_name):
    try:
        sender_email = st.secrets["email"]["account"]
        sender_password = st.secrets["email"]["password"]
    except KeyError:
        return False, "❌ 找不到 Secrets 設定！請檢查 secrets.toml。"

    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = receiver_email
    msg['Subject'] = f"[自動回報] {filename.replace('.docx', '')}"
    
    body = f"""收件人：{receiver_name}\n\n這是由系統自動生成的檢查表彙整：{filename}\n內含所有檢查項目。\n\n(由 Streamlit 雲端系統自動發送)"""
    msg.attach(MIMEText(body, 'plain'))
    part = MIMEApplication(doc_bytes, Name=filename)
    part['Content-Disposition'] = f'attachment; filename="{filename}"'
    msg.attach(part)
    
    try:
        server = smtplib.SMTP_SSL('smtp.gmail.com', 465)
        server.login(sender_email, sender_password)
        server.send_message(msg)
        server.quit()
        return True, f"✅ 寄送成功！已寄給 {receiver_name} ({receiver_email})"
    except Exception as e:
        return False, f"❌ 寄送失敗: {str(e)}"

def fetch_google_sheets_db(csv_url):
    try:
        df = pd.read_csv(csv_url)
        df = df.fillna("")
        
        required_cols = ["分類", "說明", "設計", "實測"]
        for col in required_cols:
            if col not in df.columns:
                return False, f"表單缺少必填欄位：{col}"
        
        new_db = {}
        current_category = "未分類項目"
        
        for _, row in df.iterrows():
            cat_val = str(row["分類"]).strip()
            if cat_val:
                current_category = cat_val
                
            desc = str(row["說明"]).strip()
            design = str(row["設計"]).strip()
            result = str(row["實測"]).strip()
            
            if not desc:
                continue 
            
            if current_category not in new_db:
                new_db[current_category] = []
                
            new_db[current_category].append({
                "desc": desc,
                "design": design,
                "result": result
            })
            
        return True, new_db
    except Exception as e:
        return False, f"讀取失敗：{str(e)}"

# --- 狀態管理函式 ---
def init_group_photos(g_idx):
    if f"photos_{g_idx}" not in st.session_state:
        st.session_state[f"photos_{g_idx}"] = []

def add_new_photos(g_idx, uploaded_files):
    init_group_photos(g_idx)
    current_list = st.session_state[f"photos_{g_idx}"]
    existing_ids = {p['id'] for p in current_list}
    
    for f in uploaded_files:
        file_id = f"{f.name}_{f.size}"
        if file_id not in existing_ids:
            current_list.append({
                "id": file_id, "file": f, "desc": "", "design": "", "result": "", "selected_opt_index": 0 
            })
            existing_ids.add(file_id)

def move_photo(g_idx, index, direction):
    lst = st.session_state[f"photos_{g_idx}"]
    new_index = index + direction
    if 0 <= new_index < len(lst):
        lst[index], lst[new_index] = lst[new_index], lst[index]

def delete_photo(g_idx, index):
    lst = st.session_state[f"photos_{g_idx}"]
    if 0 <= index < len(lst):
        del lst[index]

# ==========================================
# 2. 備用資料庫與常數設定
# ==========================================

RECIPIENTS = {
    "范嘉文": "ses543212004@fengyu.com.tw",
    "林憲睿": "dennys871022@fengyu.com.tw",
    "翁育玟": "Vicky1019@fengyu.com.tw",
    "林智捷": "ccl20010218@fengyu.com.tw",
    "趙健鈞": "kk919472770@fengyu.com.tw",
    "孫永明": "kevin891023@fengyu.com.tw",
    "林泓鈺": "henry30817@fengyu.com.tw",
    "黃元杰": "s10411097@fengyu.com.tw",
    "郭登慶": "tw850502@fengyu.com.tw",
    "歐冠廷": "canon1220@fengyu.com.tw",
    "黃彥榤": "ajh73684@fengyu.com.tw",
    "陳昱勳": "x85082399@fengyu.com.tw",
    "測試用 (寄給自己)": st.secrets["email"]["account"] if "email" in st.secrets else "test@example.com"
}

COMMON_SUB_CONTRACTORS = [
    "川峻工程有限公司",
    "世銓營造股份有限公司",
    "互國企業有限公司",
    "世和金屬股份有限公司",
    "宥辰興業股份有限公司",
    "亞東預拌混凝土股份有限公司",
    "自行輸入..." 
]

DEFAULT_CHECKS_DB = {
    "預設資料 (雲端連結失敗時顯示)": [
        {"desc": "這是一個預設項目", "design": "設定範例", "result": "實測範例"}
    ]
}

# ==========================================
# 3. 主程式介面邏輯
# ==========================================

st.set_page_config(page_title="工程自主檢查表生成器", layout="wide")
st.title("🏗️ 工程自主檢查表 (主控同步雲端版)")

def load_latest_db():
    if GOOGLE_SHEETS_CSV_URL.strip():
        success, result = fetch_google_sheets_db(GOOGLE_SHEETS_CSV_URL.strip())
        if success:
            return result
        else:
            st.error(f"雲端資料庫載入失敗：{result} (退回預設資料)")
            return DEFAULT_CHECKS_DB
    return DEFAULT_CHECKS_DB

if 'checks_db' not in st.session_state:
    st.session_state['checks_db'] = load_latest_db()

# Init Variables
if 'merged_doc_buffer' not in st.session_state: st.session_state['merged_doc_buffer'] = None
if 'merged_filename' not in st.session_state: st.session_state['merged_filename'] = ""
if 'saved_template' not in st.session_state: st.session_state['saved_template'] = None
if 'num_groups' not in st.session_state: st.session_state['num_groups'] = 1

DEFAULT_TEMPLATE_PATH = "template.docx"
if not st.session_state['saved_template'] and os.path.exists(DEFAULT_TEMPLATE_PATH):
    with open(DEFAULT_TEMPLATE_PATH, "rb") as f:
        st.session_state['saved_template'] = f.read()

# ==========================================
# ★ 連動邏輯：文字修改後自動推播給其他組
# ==========================================
def on_item_0_change():
    if "item_0" in st.session_state:
        base_name = st.session_state["item_0"]
        # 移除原本尾部的 #1，確保乾淨的主檔名
        if base_name.endswith("#1"):
            base_name = base_name[:-2].strip()
            
        num = st.session_state.get('num_groups', 1)
        for other_g in range(1, num):
            # 自動推播並加上自己的編號
            st.session_state[f"item_{other_g}"] = f"{base_name}#{other_g + 1}"

def update_group_info(g_idx):
    base_date = st.session_state.get('global_date', datetime.date.today())
    selected_type = st.session_state[f"type_{g_idx}"]
    item_name, _ = generate_names(selected_type, base_date)
    
    # ★ 自動加上 #1, #2 等編號
    st.session_state[f"item_{g_idx}"] = f"{item_name}      #{g_idx + 1}"
    
    def clear_group_data(idx):
        keys_to_clear = [k for k in st.session_state.keys() if f"_{idx}_" in k and (k.startswith("sel_") or k.startswith("desc_") or k.startswith("design_") or k.startswith("result_"))]
        for k in keys_to_clear: del st.session_state[k]
        if f"photos_{idx}" in st.session_state:
            for p in st.session_state[f"photos_{idx}"]:
                p['desc'] = ""; p['design'] = ""; p['result'] = ""; p['selected_opt_index'] = 0

    clear_group_data(g_idx)
    
    if g_idx == 0:
        current_num_groups = st.session_state.get('num_groups', 1)
        for other_g in range(1, current_num_groups):
            st.session_state[f"type_{other_g}"] = selected_type
            # 確保同步時，其他組別擁有正確的 #2, #3 編號
            st.session_state[f"item_{other_g}"] = f"{item_name}#{other_g + 1}"
            clear_group_data(other_g)

def clear_all_data():
    for key in list(st.session_state.keys()):
        if key.startswith(('type_', 'item_', 'fname_', 'photos_', 'file_', 'sel_', 'desc_', 'design_', 'result_')):
            del st.session_state[key]
    st.session_state['num_groups'] = 1
    st.session_state['merged_doc_buffer'] = None
    st.session_state['merged_filename'] = ""

# Sidebar
with st.sidebar:
    st.header("1. 樣板設定")
    if st.session_state['saved_template']:
        st.success("✅ Word 樣板已載入")
    else:
        uploaded = st.file_uploader("上傳樣板", type=['docx'])
        if uploaded:
            st.session_state['saved_template'] = uploaded.getvalue()
            st.rerun()
            
    st.markdown("---")
    st.header("☁️ 雲端資料庫狀態")
    if GOOGLE_SHEETS_CSV_URL.strip():
        st.success("✅ 已綁定專屬試算表")
        if st.button("🔄 點我強制同步最新資料", use_container_width=True, type="primary"):
            with st.spinner("📥 正在抓取最新資料..."):
                st.session_state['checks_db'] = load_latest_db()
                st.success("更新完成！")
                st.rerun()
    else:
        st.warning("⚠️ 尚未設定 GOOGLE_SHEETS_CSV_URL。")
            
    st.markdown("---")
    st.button("🗑️ 清除所有填寫資料", on_click=clear_all_data, use_container_width=True)

    st.markdown("---")
    st.header("2. 專案資訊")
    p_name = st.text_input("工程名稱", "衛生福利部防疫中心興建工程")
    p_cont = st.text_input("施工廠商", "豐譽營造股份有限公司")
    sub_select = st.selectbox("協力廠商", COMMON_SUB_CONTRACTORS)
    if sub_select == "自行輸入...":
        p_sub = st.text_input("請輸入廠商名稱", "川峻工程有限公司")
    else:
        p_sub = sub_select
    p_loc = st.text_input("施作位置", "北棟 1F")
    base_date = st.date_input("日期", get_taiwan_date(), key='global_date')

# Main Body
if st.session_state['saved_template']:
    num_groups = st.number_input("本次產生幾組檢查表？", min_value=1, value=st.session_state['num_groups'], key='num_groups_input')
    st.session_state['num_groups'] = num_groups
    all_groups_data = []

    for g in range(num_groups):
        st.markdown(f"---")
        st.subheader(f"📂 第 {g+1} 組")
        c1, c2, c3 = st.columns([2, 2, 1])
        db_options = list(st.session_state['checks_db'].keys())
        
        # ==========================================
        # ★ 剛新增組別時，自動預設帶入第一組的選項及名稱
        # ==========================================
        if g > 0 and f"type_{g}" not in st.session_state and "type_0" in st.session_state:
            st.session_state[f"type_{g}"] = st.session_state["type_0"]
            if "item_0" in st.session_state:
                base_name = st.session_state["item_0"]
                if base_name.endswith("#1"):
                    base_name = base_name[:-2].strip()
                st.session_state[f"item_{g}"] = f"{base_name}#{g + 1}"
            
        selected_type = c1.selectbox(f"選擇檢查工項", db_options, key=f"type_{g}", on_change=update_group_info, args=(g,))
        
        if f"item_{g}" not in st.session_state:
            update_group_info(g)
            
        # ★ 第一組綁定 on_change，只要修改就會同步全場
        if g == 0:
            g_item = c2.text_input(f"自檢項目名稱", key=f"item_{g}", on_change=on_item_0_change)
        else:
            g_item = c2.text_input(f"自檢項目名稱", key=f"item_{g}")
            
        roc_year = base_date.year - 1911
        date_display = f"{roc_year}.{base_date.month:02d}.{base_date.day:02d}"
        c3.text(f"日期: {date_display}")

        st.markdown("##### 📸 照片上傳與排序")
        uploader_key_name = f"uploader_key_{g}"
        if uploader_key_name not in st.session_state: st.session_state[uploader_key_name] = 0
        dynamic_key = f"uploader_{g}_{st.session_state[uploader_key_name]}"
        
        new_files = st.file_uploader(f"點擊此處選擇照片 (第 {g+1} 組)", type=['jpg','png','jpeg'], accept_multiple_files=True, key=dynamic_key)
        if new_files:
            add_new_photos(g, new_files)
            st.session_state[uploader_key_name] += 1
            st.rerun()
        
        if st.session_state.get(f"photos_{g}"):
            if st.button("🔄 順序反了嗎？點我「一鍵反轉」照片順序", key=f"rev_{g}"):
                current_list = st.session_state[f"photos_{g}"]
                for p in current_list:
                    d_key = f"desc_{g}_{p['id']}"
                    if d_key in st.session_state: p['desc'] = st.session_state[d_key]
                    des_key = f"design_{g}_{p['id']}"
                    if des_key in st.session_state: p['design'] = st.session_state[des_key]
                    r_key = f"result_{g}_{p['id']}"
                    if r_key in st.session_state: p['result'] = st.session_state[r_key]
                    s_key = f"sel_{g}_{p['id']}"
                    if s_key in st.session_state: p['selected_opt_index'] = st.session_state[s_key]
                st.session_state[f"photos_{g}"].reverse()
                st.rerun()
        
        init_group_photos(g)
        photo_list = st.session_state[f"photos_{g}"]
        
        if photo_list:
            check_items_list = st.session_state['checks_db'].get(selected_type, [])
            options = ["(請選擇...)"] + [item['desc'] for item in check_items_list]

            for i, photo_data in enumerate(photo_list):
                with st.container():
                    col_img, col_info, col_ctrl = st.columns([1.5, 3, 0.5])
                    pid = photo_data['id']
                    with col_img:
                        st.image(photo_data['file'], use_container_width=True)
                        st.caption(f"No. {i+1:02d}")
                    with col_info:
                        def on_select_change(pk=pid, gk=g):
                            k = f"sel_{gk}_{pk}"
                            if k not in st.session_state: return
                            new_idx = st.session_state[k]
                            dk, desk, rk = f"desc_{gk}_{pk}", f"design_{gk}_{pk}", f"result_{gk}_{pk}"
                            if isinstance(new_idx, int) and new_idx > 0 and new_idx <= len(check_items_list):
                                item_data = check_items_list[new_idx-1]
                                st.session_state[dk] = item_data['desc']
                                st.session_state[desk] = item_data['design']
                                st.session_state[rk] = item_data['result']
                            else:
                                st.session_state[dk] = ""
                                st.session_state[desk] = ""
                                st.session_state[rk] = ""

                        current_opt_idx = photo_data.get('selected_opt_index', 0)
                        if current_opt_idx > len(options): current_opt_idx = 0
                        st.selectbox("快速填寫", range(len(options)), format_func=lambda x: options[x], index=current_opt_idx, key=f"sel_{g}_{pid}", on_change=on_select_change, label_visibility="collapsed")

                        def on_text_change(field, pk=pid, idx=i, gk=g): 
                            val = st.session_state[f"{field}_{gk}_{pk}"]
                            st.session_state[f"photos_{gk}"][idx][field] = val
                            if field == 'sel': st.session_state[f"photos_{gk}"][idx]['selected_opt_index'] = val

                        desc_key = f"desc_{g}_{pid}"
                        design_key = f"design_{g}_{pid}"
                        result_key = f"result_{g}_{pid}"
                        if desc_key not in st.session_state: st.session_state[desc_key] = photo_data.get('desc', '')
                        if design_key not in st.session_state: st.session_state[design_key] = photo_data.get('design', '')
                        if result_key not in st.session_state: st.session_state[result_key] = photo_data.get('result', '')

                        st.text_input("說明", key=desc_key, on_change=on_text_change, args=('desc',))
                        st.text_input("設計 (可留空)", key=design_key, on_change=on_text_change, args=('design',))
                        st.text_input("實測", key=result_key, on_change=on_text_change, args=('result',))

                    with col_ctrl:
                        if st.button("⬆️", key=f"up_{g}_{i}"): move_photo(g, i, -1); st.rerun()
                        if st.button("⬇️", key=f"down_{g}_{i}"): move_photo(g, i, 1); st.rerun()
                        if st.button("❌", key=f"del_{g}_{i}"): delete_photo(g, i); st.rerun()
                    st.divider()

            g_photos_export = []
            for i, p in enumerate(photo_list):
                d_val = st.session_state.get(f"desc_{g}_{p['id']}", p['desc'])
                des_val = st.session_state.get(f"design_{g}_{p['id']}", p['design'])
                r_val = st.session_state.get(f"result_{g}_{p['id']}", p['result'])
                g_photos_export.append({
                    "file": p['file'], "no": i + 1, "date_str": date_display, 
                    "desc": d_val, "design": des_val, "result": r_val
                })

            all_groups_data.append({
                "group_id": g+1,
                "context": {
                    "project_name": p_name, "contractor": p_cont, "sub_contractor": p_sub,
                    "location": p_loc, "date": date_display, "check_item": g_item
                },
                "photos": g_photos_export
            })

    st.markdown("---")
    st.subheader("🚀 執行操作")
    default_filename = ""
    if "type_0" in st.session_state:
        default_filename = generate_clean_filename_base(st.session_state["type_0"], base_date)
    else:
        default_filename = f"自主檢查表_{get_taiwan_date()}"

    final_file_name_input = st.text_input("📝 最終 Word 檔名", value=default_filename)
    if not final_file_name_input.endswith(".docx"): final_file_name = final_file_name_input + ".docx"
    else: final_file_name = final_file_name_input

    selected_name = st.selectbox("📬 收件人", list(RECIPIENTS.keys()))
    target_email = RECIPIENTS[selected_name]

    if st.button("步驟 1：生成報告資料 (單一 Word 檔)", type="primary", use_container_width=True):
        if not all_groups_data: st.error("⚠️ 請至少上傳一張照片並填寫資料")
        else:
            with st.spinner("📦 正在生成並合併 Word 檔案..."):
                master_doc = None
                composer = None
                for group in all_groups_data:
                    photos = group['photos']
                    context = group['context']
                    for page_idx, i in enumerate(range(0, len(photos), 8)):
                        batch = photos[i : i+8]
                        start_no = i + 1
                        current_doc = generate_single_page(st.session_state['saved_template'], context, batch, start_no)
                        if master_doc is None:
                            master_doc = current_doc
                            composer = Composer(master_doc)
                        else:
                            composer.append(current_doc)
                out_buffer = io.BytesIO()
                composer.save(out_buffer)
                st.session_state['merged_doc_buffer'] = out_buffer.getvalue()
                st.session_state['merged_filename'] = final_file_name
                st.success(f"✅ 彙整完成！檔名：{final_file_name}")

    if st.session_state['merged_doc_buffer']:
        col_mail, col_dl = st.columns(2)
        with col_mail:
            if st.button(f"📧 立即寄出 Word 檔給：{selected_name}", use_container_width=True):
                with st.spinner("📨 雲端發信中..."):
                    success, msg = send_email_via_secrets(st.session_state['merged_doc_buffer'], st.session_state['merged_filename'], target_email, selected_name)
                    if success: st.success(msg)
                    else: st.error(msg)
        with col_dl:
            st.download_button(label="📥 下載 Word 檔案", data=st.session_state['merged_doc_buffer'], file_name=st.session_state['merged_filename'], mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
else:
    st.info("👈 請先在左側確認 Word 樣板")
